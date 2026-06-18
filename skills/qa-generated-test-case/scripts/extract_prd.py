#!/usr/bin/env python3
"""Extract PRD text from Word, Markdown, or text files for CaseForge."""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def clean_text(value: str) -> str:
    return "\n".join(line.strip() for line in value.splitlines() if line.strip()).strip()


def extract_with_python_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("python-docx unavailable") from exc

    doc = Document(path)
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table_index, table in enumerate(doc.tables, 1):
        rows: list[str] = []
        for row in table.rows:
            cells = [clean_text(cell.text).replace("\n", " / ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            lines.append(f"[表格 {table_index}]")
            lines.extend(rows)

    return clean_text("\n".join(lines))


def element_text(element: ET.Element) -> str:
    return "".join(node.text or "" for node in element.findall(".//w:t", WORD_NS)).strip()


def extract_docx_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))

    body = root.find("w:body", WORD_NS)
    if body is None:
        return ""

    lines: list[str] = []
    table_index = 0
    for child in body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = element_text(child)
            if text:
                lines.append(text)
        elif tag == "tbl":
            table_index += 1
            lines.append(f"[表格 {table_index}]")
            for row in child.findall(".//w:tr", WORD_NS):
                cells = [element_text(cell).replace("\n", " / ") for cell in row.findall("./w:tc", WORD_NS)]
                if any(cells):
                    lines.append(" | ".join(cells))

    return clean_text("\n".join(lines))


def extract_with_textutil(path: Path) -> str:
    if shutil.which("textutil") is None:
        raise RuntimeError("textutil is unavailable; cannot extract this Word format")
    result = subprocess.run(
        ["textutil", "-stdout", "-convert", "txt", str(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    return clean_text(result.stdout)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return clean_text(path.read_text(encoding="utf-8"))
    if suffix == ".docx":
        for extractor in (extract_with_python_docx, extract_docx_xml):
            try:
                text = extractor(path)
                if text:
                    return text
            except Exception:
                continue
    if suffix in {".doc", ".docx"}:
        return extract_with_textutil(path)
    raise RuntimeError(f"Unsupported PRD file type: {suffix or '<none>'}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract PRD text for CaseForge.")
    parser.add_argument("input", type=Path, help="Input .docx/.doc/.txt/.md file")
    parser.add_argument("--out", type=Path, help="Output text file")
    args = parser.parse_args()

    if not args.input.exists():
        raise SystemExit(f"Input file not found: {args.input}")

    text = extract_text(args.input)
    if not text:
        raise SystemExit("No text extracted.")

    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(text, encoding="utf-8")
        print(f"out={args.out}")
    else:
        sys.stdout.write(text)
        if not text.endswith("\n"):
            sys.stdout.write("\n")
    print(f"chars={len(text)}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
